import torch
from transformers import AutoModelForSpeechSeq2Seq, AutoProcessor, pipeline
from moviepy.editor import VideoFileClip
from docx import Document
import cv2
from facenet_pytorch import MTCNN, InceptionResnetV1
import numpy as np
import random
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize, sent_tokenize
import gensim
from gensim import corpora
from gensim.models.ldamodel import LdaModel
import spacy
from textstat.textstat import textstatistics
from clothes import process_images

# Setup Whisper Model
model_id = "openai/whisper-large-v3"
device = "cuda:0" if torch.cuda.is_available() else "cpu"
torch_dtype = torch.float16 if torch.cuda.is_available() else torch.float32

model = AutoModelForSpeechSeq2Seq.from_pretrained(model_id, torch_dtype=torch_dtype, low_cpu_mem_usage=True, use_safetensors=True)
model.to(device)

processor = AutoProcessor.from_pretrained(model_id)

pipe = pipeline(
    "automatic-speech-recognition",
    model=model,
    tokenizer=processor.tokenizer,
    feature_extractor=processor.feature_extractor,
    max_new_tokens=128,
    chunk_length_s=15,
    batch_size=16,
    return_timestamps=True,
    torch_dtype=torch_dtype,
    device=device,
)

# Extract audio from video and save as MP3
video_path = "/Users/riditjain/Downloads/WhatsApp Video 2024-05-16 at 17.09.02.mp4"
audio_path = "/Users/riditjain/Downloads/WhatsApp Video 2024-05-16 at 17.09.02.mp3"
video_clip = VideoFileClip(video_path)
audio_clip = video_clip.audio
audio_clip.write_audiofile(audio_path, codec='mp3')

# Transcribe audio
result = pipe(audio_path)
transcribed_text = result['text']

# Save transcribed text to Word document
doc = Document()
doc.add_heading('Translated Text', level=1)
doc.add_paragraph(transcribed_text)

# NLP Analysis
nlp = spacy.load("en_core_web_sm")

def preprocess_text(text):
    text = text.lower()
    sentences = sent_tokenize(text)
    words = [word_tokenize(sentence) for sentence in sentences]
    stop_words = set(stopwords.words('english'))
    filtered_words = [[word for word in word_list if word not in stop_words] for word_list in words]
    return filtered_words

def identify_keywords(text):
    processed_text = preprocess_text(text)
    all_words = [word for sublist in processed_text for word in sublist]
    freq_dist = nltk.FreqDist(all_words)
    keywords = [word for word, freq in freq_dist.items() if freq > 1]
    return keywords

def topic_modeling(text):
    processed_text = preprocess_text(text)
    dictionary = corpora.Dictionary(processed_text)
    doc_term_matrix = [dictionary.doc2bow(doc) for doc in processed_text]
    lda_model = LdaModel(doc_term_matrix, num_topics=5, id2word=dictionary, passes=25)
    topics = lda_model.print_topics(num_words=3)
    return topics

def named_entity_recognition(text):
    doc = nlp(text)
    entities = [(entity.text, entity.label_) for entity in doc.ents]
    return entities

def analyze_complexity(text):
    syllables = textstatistics().syllable_count(text)
    sentences = textstatistics().sentence_count(text)
    words = len(word_tokenize(text))
    flesch_kincaid_grade = 0.39 * (words/sentences) + 11.8 * (syllables/words) - 15.59
    return flesch_kincaid_grade

def assess_technology_experience(text):
    doc = nlp(text)
    experiences = []

    for sentence in doc.sents:
        has_development_verb = False
        tech_skills = []

        for token in sentence:
            if token.dep_ == "ROOT" and token.lemma_ in {"develop", "build", "create", "implement", "use", "utilize"}:
                has_development_verb = True

            if token.dep_ in {"dobj", "pobj"} and token.pos_ == "PROPN":
                tech_skills.append(token.text)

        if has_development_verb and tech_skills:
            for skill in tech_skills:
                experiences.append(f"The candidate has intermediate to advanced experience with {skill} based on the context: '{sentence.text}'")
        elif tech_skills:
            for skill in tech_skills:
                experiences.append(f"The candidate mentioned {skill}, indicating familiarity: '{sentence.text}'")
    
    return experiences

def analyze_thought_process(text):
    sentences = sent_tokenize(text)
    logical_connectors = ['therefore', 'however', 'for example', 'because', 'firstly', 'secondly', 'furthermore', 'moreover']
    connector_count = 0
    sentence_lengths = []

    for sentence in sentences:
        words = word_tokenize(sentence)
        sentence_lengths.append(len(words))
        connector_count += sum(1 for word in words if word.lower() in logical_connectors)

    avg_sentence_length = sum(sentence_lengths) / len(sentence_lengths) if sentence_lengths else 0

    return {
        'total_sentences': len(sentences),
        'average_sentence_length': avg_sentence_length,
        'logical_connector_count': connector_count,
    }

thought_process_metrics = analyze_thought_process(transcribed_text)
keywords = identify_keywords(transcribed_text)
topics = topic_modeling(transcribed_text)
entities = named_entity_recognition(transcribed_text)
complexity_score = analyze_complexity(transcribed_text)
technology_experience = assess_technology_experience(transcribed_text)

doc.add_heading('NLP Results', level=1)
doc.add_heading('Thought Process Metrics', level=2)
doc.add_paragraph(f"Total Sentences: {thought_process_metrics['total_sentences']}")
doc.add_paragraph(f"Average Sentence Length: {thought_process_metrics['average_sentence_length']:.2f}")
doc.add_paragraph(f"Logical Connector Count: {thought_process_metrics['logical_connector_count']}")

doc.add_heading('Keywords', level=2)
doc.add_paragraph(", ".join(keywords))

doc.add_heading('Topics', level=2)
for topic in topics:
    doc.add_paragraph(str(topic))

doc.add_heading('Named Entities', level=2)
for entity in entities:
    doc.add_paragraph(f"{entity[0]}: {entity[1]}")

doc.add_heading('Complexity Score', level=2)
doc.add_paragraph(str(complexity_score))

doc.add_heading('Technology Experience', level=2)
for experience in technology_experience:
    doc.add_paragraph(experience)

# Face Detection
device = torch.device('cuda:0' if torch.cuda.is_available() else 'cpu')
mtcnn = MTCNN(keep_all=True, device=device)
resnet = InceptionResnetV1(pretrained='vggface2').eval().to(device)

video_capture = cv2.VideoCapture(video_path)

persistent_embeddings = {}
next_id = 1
cheating_log = []

def find_matching_id(new_embedding, embeddings, threshold=0.8):
    if len(embeddings) == 0:
        return None
    min_distance = float('inf')
    matching_id = None
    for face_id, emb in embeddings.items():
        distance = torch.nn.functional.pairwise_distance(new_embedding, emb).min().item()
        if distance < min_distance:
            min_distance = distance
            matching_id = face_id
    if min_distance < threshold:
        return matching_id
    return None

while True:
    ret, frame = video_capture.read()
    if not ret:
        break

    rgb_frame = cv2.cvtColor(frame, cv2.COLOR_BGR2RGB)
    boxes, _ = mtcnn.detect(rgb_frame)

    if boxes is not None:
        for box in boxes:
            x1, y1, x2, y2 = [max(0, int(coord)) for coord in box]
            x2 = min(x2, rgb_frame.shape[1])
            y2 = min(y2, rgb_frame.shape[0])
            
            face = rgb_frame[y1:y2, x1:x2]
            if face.size == 0:
                continue
            
            face = cv2.resize(face, (160, 160))
            face = torch.tensor(face).permute(2,0,1).float().div(255).unsqueeze(0).to(device)
            
            with torch.no_grad():
                embedding = resnet(face).detach()
            
            matching_id = find_matching_id(embedding, persistent_embeddings)

            if matching_id is None:
                persistent_embeddings[next_id] = embedding
                cheating_log.append(f"Person {next_id} enters.")
                next_id += 1
            else:
                cheating_log.append(f"Person {matching_id} re-enters.")

video_capture.release()
cv2.destroyAllWindows()

doc.add_heading('Cheating', level=1)
for log_entry in cheating_log:
    doc.add_paragraph(log_entry)

# Extract 5 random frames and analyze clothing
video_capture = cv2.VideoCapture(video_path)
frame_count = int(video_capture.get(cv2.CAP_PROP_FRAME_COUNT))
random_frames = random.sample(range(frame_count), 5)

clothing_log = []

frame_paths = []

for frame_no in random_frames:
    video_capture.set(cv2.CAP_PROP_POS_FRAMES, frame_no)
    ret, frame = video_capture.read()
    if ret:
        image_path = f"/Users/riditjain/Downloads/frame_{frame_no}.jpg"
        cv2.imwrite(image_path, frame)
        frame_paths.append(image_path)

video_capture.release()

clothing_results = process_images(frame_paths, threshold=0.8, show_image=True)

for frame_no, result in zip(random_frames, clothing_results):
    clothing_log.append(f"Frame {frame_no}: {result}")

doc.add_heading('Clothing', level=1)
for log_entry in clothing_log:
    doc.add_paragraph(log_entry)

# Save the document
doc.save('/Users/riditjain/Downloads/InterviewProFinal-main 2/logfile2.docx')
 